#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一文档提取器 - 支持 PDF、Word、Markdown、RTF
依赖：pip install PyMuPDF python-docx
"""

import argparse
import json
import sys
from pathlib import Path


def extract_pdf(file_path: str) -> dict:
    """提取 PDF 文档内容"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("错误：请安装 PyMuPDF: pip install PyMuPDF", file=sys.stderr)
        sys.exit(1)

    doc = fitz.open(file_path)
    content = {
        "format": "pdf",
        "source": file_path,
        "pages": [],
        "full_text": ""
    }

    full_text = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        content["pages"].append({
            "page": page_num,
            "text": text
        })
        full_text.append(text)

    content["full_text"] = "\n".join(full_text)
    doc.close()
    return content


def extract_docx(file_path: str) -> dict:
    """提取 Word 文档内容"""
    try:
        from docx import Document
    except ImportError:
        print("错误：请安装 python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(file_path)
    content = {
        "format": "docx",
        "source": file_path,
        "paragraphs": [],
        "tables": [],
        "full_text": ""
    }

    # 提取段落
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else None
            })
    content["paragraphs"] = paragraphs

    # 提取表格
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append({
            "index": table_idx,
            "data": table_data
        })

    content["full_text"] = "\n".join([p["text"] for p in paragraphs])
    return content


def extract_markdown(file_path: str) -> dict:
    """提取 Markdown 文档内容"""
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    return {
        "format": "markdown",
        "source": file_path,
        "full_text": text
    }


def extract_document(file_path: str, format_hint: str = None) -> dict:
    """根据格式提取文档"""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 自动检测格式
    suffix = path.suffix.lower()
    if format_hint:
        format_type = format_hint.lower()
    elif suffix == '.pdf':
        format_type = 'pdf'
    elif suffix in ['.docx', '.doc']:
        format_type = 'docx'
    elif suffix in ['.md', '.markdown']:
        format_type = 'markdown'
    else:
        format_type = 'text'

    # 提取内容
    if format_type == 'pdf':
        return extract_pdf(file_path)
    elif format_type == 'docx':
        return extract_docx(file_path)
    elif format_type in ['markdown', 'md']:
        return extract_markdown(file_path)
    else:
        # 默认作为纯文本处理
        with open(file_path, 'r', encoding='utf-8') as f:
            return {
                "format": "text",
                "source": file_path,
                "full_text": f.read()
            }


def main():
    parser = argparse.ArgumentParser(description='提取文档内容（PDF/Word/Markdown）')
    parser.add_argument('-i', '--input', required=True, help='输入文件路径')
    parser.add_argument('-f', '--format', help='文档格式（pdf/docx/md），不指定则自动检测')
    parser.add_argument('-o', '--output', help='输出 JSON 文件路径，不指定则输出到 stdout')
    args = parser.parse_args()

    try:
        content = extract_document(args.input, args.format)

        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(content, f, ensure_ascii=False, indent=2)
            print(f"已提取到: {args.output}")
        else:
            print(json.dumps(content, ensure_ascii=False, indent=2))

    except Exception as e:
        print(f"提取失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
